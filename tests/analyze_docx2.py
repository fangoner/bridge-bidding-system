from docx import Document

doc = Document(r'd:\Bridge Card\Bidding System\JF实战_标准自然_For DIFY - Rev 3.1.docx')

print(f"总段落数: {len(doc.paragraphs)}")
print("=" * 80)

for i, para in enumerate(doc.paragraphs[100:200]):
    text = para.text.strip()
    if text:
        style_name = para.style.name if para.style else "None"
        preview = text[:150] + "..." if len(text) > 150 else text
        print(f"[{i+100}] {style_name}")
        print(f"    {preview}")
        print()
