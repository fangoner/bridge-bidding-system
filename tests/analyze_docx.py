from docx import Document

doc = Document(r'd:\Bridge Card\Bidding System\JF实战_标准自然_For DIFY - Rev 3.1.docx')

print(f"总段落数: {len(doc.paragraphs)}")
print("=" * 80)

for i, para in enumerate(doc.paragraphs[:100]):
    text = para.text.strip()
    if text:
        style_name = para.style.name if para.style else "None"
        indent = 0
        if para.paragraph_format.left_indent:
            indent = int(para.paragraph_format.left_indent / 360)  # 转换为缩进级别
        
        preview = text[:100] + "..." if len(text) > 100 else text
        print(f"[{i}] Style:{style_name}, Indent:{indent}")
        print(f"    {preview}")
        print()
